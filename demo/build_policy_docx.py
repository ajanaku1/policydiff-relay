"""Build the Google Docs-ready Eligibility Policy demo document."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLACK = RGBColor(0x00, 0x00, 0x00)
MUTED = RGBColor(0x55, 0x55, 0x55)
FONT = "Arial"


def set_font_family(font, run_properties) -> None:
    font.name = FONT
    run_properties.rFonts.set(qn("w:ascii"), FONT)
    run_properties.rFonts.set(qn("w:hAnsi"), FONT)


def set_run_font(run, size: float, color=BLACK, bold: bool = False) -> None:
    set_font_family(run.font, run._element.get_or_add_rPr())
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold


def set_paragraph_spacing(target, before: float, after: float, line: float) -> None:
    formatting = getattr(target, "paragraph_format", target)
    formatting.space_before = Pt(before)
    formatting.space_after = Pt(after)
    formatting.line_spacing = line
    formatting.widow_control = True


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_normal_style(document: Document) -> None:
    style = document.styles["Normal"]
    set_font_family(style.font, style._element.get_or_add_rPr())
    style.font.size = Pt(11)
    style.font.color.rgb = BLACK
    set_paragraph_spacing(style.paragraph_format, 0, 8, 1.15)


def configure_heading(style, size: float, color, before: float, after: float) -> None:
    set_font_family(style.font, style._element.get_or_add_rPr())
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = False
    set_paragraph_spacing(style.paragraph_format, before, after, 1.15)
    style.paragraph_format.keep_with_next = True


def configure_styles(document: Document) -> None:
    configure_normal_style(document)
    configure_heading(document.styles["Heading 1"], 20, BLACK, 20, 6)
    configure_heading(document.styles["Heading 2"], 16, BLACK, 18, 6)
    configure_heading(document.styles["Heading 3"], 14, RGBColor(0x43, 0x43, 0x43), 16, 4)


def add_title_block(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(title, 0, 3, 1.0)
    title.paragraph_format.keep_with_next = True
    set_run_font(title.add_run("Eligibility Policy"), 26)

    metadata = document.add_paragraph()
    set_paragraph_spacing(metadata, 0, 10, 1.15)
    metadata.paragraph_format.keep_with_next = True
    set_run_font(
        metadata.add_run("Northstar Benefits  ·  Version 4  ·  Effective June 1, 2026"),
        11,
        MUTED,
    )


def add_clause(document: Document, heading: str, body: str) -> None:
    document.add_paragraph(heading, style="Heading 2")
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_together = True
    set_run_font(paragraph.add_run(body), 11)


def build_document(output_path: Path) -> None:
    document = Document()
    configure_page(document)
    configure_styles(document)
    add_title_block(document)
    add_clause(
        document,
        "§4.2 Eligibility age",
        "Applicants must be at least 18 years old on the date of enrollment.",
    )
    add_clause(
        document,
        "§7.1 Waiting period",
        "The 30-day waiting period begins on the confirmed application date.",
    )
    add_clause(
        document,
        "§9.3 Contractor classification",
        "Contractor eligibility follows the applicable written agreement.",
    )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


if __name__ == "__main__":
    build_document(Path(__file__).with_name("Eligibility Policy.raw.docx"))
